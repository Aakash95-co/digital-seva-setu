import os
import base64
import datetime
import json
import io

import dash
from dash import html, dcc, Input, Output, State, ALL, callback_context
import dash_bootstrap_components as dbc
import pandas as pd

from app import app
from data import df_adv, FY_DATA

# ─────────────────────────────────────────────────────────────────────────────
# UPLOAD FOLDER
# ─────────────────────────────────────────────────────────────────────────────
_UPLOAD_DIR = os.path.normpath(
    os.path.join(os.path.dirname(__file__), '..', 'uploads')
)
os.makedirs(_UPLOAD_DIR, exist_ok=True)

# ─────────────────────────────────────────────────────────────────────────────
# DATA PREP  (dropdown options)
# ─────────────────────────────────────────────────────────────────────────────
def _prep_options(raw):
    if raw is None or raw.empty:
        return [], [], []
    df = raw.copy()
    df.columns = df.columns.str.strip()
    _map = {
        'District_name': 'District', 'District_Eng': 'District',
        'Office_name': 'Office',     'Office_Eng': 'Office',
        'Service_name': 'Service',   'Service_Eng': 'Service',
    }
    df.rename(columns={k: v for k, v in _map.items() if k in df.columns}, inplace=True)
    for c in ('District', 'Office', 'Service'):
        if c in df.columns:
            df[c] = df[c].astype(str).str.strip()
    districts = sorted(df['District'].dropna().unique()) if 'District' in df.columns else []
    offices   = sorted(df['Office'].dropna().unique())   if 'Office'   in df.columns else []
    services  = sorted(df['Service'].dropna().unique())  if 'Service'  in df.columns else []
    return districts, offices, services


_districts, _offices, _services = _prep_options(df_adv)
_district_opts = [{'label': d, 'value': d} for d in _districts]
_office_opts   = [{'label': o, 'value': o} for o in _offices]
_service_opts  = [{'label': s, 'value': s} for s in _services]

_ICON_MAP = {
    'PDF': '📄', 'XLSX': '📊', 'XLS': '📊', 'CSV': '📋',
    'DOCX': '📝', 'DOC': '📝', 'JPG': '🖼️', 'JPEG': '🖼️',
    'PNG': '🖼️', 'ZIP': '🗜️',
}

def _entry_key(e):
    return f"D:{e.get('district','')};O:{e.get('office','')};S:{e.get('service','')}"


# ─────────────────────────────────────────────────────────────────────────────
# LAYOUT
# ─────────────────────────────────────────────────────────────────────────────
layout = html.Div([
    # Header
    html.Div([
        html.H2("📁 Report & Tracking",
                style={'color': 'white', 'margin': '0', 'fontSize': '1.6rem'}),
        html.P("Upload reports for record-keeping  |  Track districts, offices or services over time",
               style={'color': '#c8dff0', 'margin': '4px 0 0 0', 'fontSize': '0.9rem'}),
    ], style={'background': 'linear-gradient(90deg,#1a3c5e,#2d6a9f)',
              'padding': '18px 28px', 'borderRadius': '10px', 'marginBottom': '24px'}),

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 1 — REPORTING
    # ══════════════════════════════════════════════════════════════════════
    html.Div([
        html.H4("📤 Reporting — Upload Files",
                style={'color': '#1a3c5e', 'borderBottom': '2px solid #2d6a9f',
                       'paddingBottom': '8px', 'marginBottom': '16px'}),
        html.P("Upload reports (PDF, Excel, CSV, Word, images …). "
               "Files are saved on the server and listed below for future reference.",
               style={'color': '#666', 'fontSize': '0.92rem', 'marginBottom': '14px'}),

        # Remarks input
        html.Div([
            html.Label("📝 Remarks (optional)",
                       style={'fontWeight': 'bold', 'marginBottom': '6px', 'color': '#1a3c5e'}),
            dcc.Textarea(
                id='rt-upload-remark',
                placeholder='Add a remark or description for the file(s) being uploaded…',
                style={
                    'width': '100%', 'height': '70px', 'borderRadius': '6px',
                    'border': '1px solid #b0cce4', 'padding': '8px 12px',
                    'fontSize': '0.92rem', 'resize': 'vertical',
                    'fontFamily': 'inherit',
                },
            ),
        ], style={'marginBottom': '14px'}),

        dcc.Upload(
            id='rt-upload',
            children=html.Div([
                html.Div("📂", style={'fontSize': '2.5rem', 'marginBottom': '8px'}),
                html.Strong("Drag & Drop files here"),
                html.Span(" or ", style={'color': '#888'}),
                html.A("Browse", style={'color': '#2d6a9f', 'textDecoration': 'underline',
                                        'cursor': 'pointer'}),
                html.Br(),
                html.Small("Supported: PDF, XLSX, XLS, CSV, DOCX, JPG, PNG, …",
                           style={'color': '#888'}),
            ], style={'textAlign': 'center', 'padding': '10px'}),
            style={
                'width': '100%', 'minHeight': '130px', 'lineHeight': '1.6',
                'borderWidth': '2px', 'borderStyle': 'dashed', 'borderColor': '#2d6a9f',
                'borderRadius': '10px', 'backgroundColor': '#f8fbff',
                'display': 'flex', 'alignItems': 'center',
                'justifyContent': 'center', 'cursor': 'pointer',
            },
            multiple=True,
        ),
        html.Div(id='rt-upload-status', style={'marginTop': '10px'}),

        # Saved files table
        html.Div([
            html.H6("📋 Uploaded Files",
                    style={'color': '#1a3c5e', 'fontWeight': 'bold', 'marginBottom': '10px'}),
            html.Div(id='rt-uploads-list'),
        ], style={'background': '#f8fbff', 'border': '1px solid #d0e4f4',
                  'borderRadius': '8px', 'padding': '14px 18px', 'marginTop': '16px'}),

    ], style={
        'background': 'white', 'border': '1px solid #d0d7de',
        'borderRadius': '10px', 'padding': '22px 26px', 'marginBottom': '20px',
        'boxShadow': '0 2px 8px rgba(0,0,0,0.06)',
    }),

    # ══════════════════════════════════════════════════════════════════════
    # REPORT GENERATION BUTTONS (after Reporting section)
    # ══════════════════════════════════════════════════════════════════════
    html.Div([
        html.H5("📊 Generate Reports", style={
            'color': '#1a3c5e', 'marginBottom': '14px', 'fontWeight': 'bold',
        }),
        dbc.Row([
            dbc.Col([
                dbc.Button([
                    html.Span("🤖", style={'fontSize': '1.4rem', 'marginRight': '8px',
                                           'verticalAlign': 'middle'}),
                    html.Span([
                        html.Strong("Generate AI Report", style={'display': 'block',
                                                                  'fontSize': '1rem'}),
                        html.Small("District & Service Summary  (Word .docx)",
                                   style={'color': '#c8dff0', 'fontWeight': 'normal',
                                          'fontSize': '0.82rem'}),
                    ], style={'display': 'inline-block', 'verticalAlign': 'middle'}),
                ],
                    id='rt-ai-report-btn',
                    style={
                        'width': '100%', 'padding': '14px 20px',
                        'borderRadius': '10px', 'textAlign': 'left',
                        'background': 'linear-gradient(135deg,#1a3c5e,#2d6a9f)',
                        'border': 'none', 'boxShadow': '0 3px 10px rgba(45,106,159,0.4)',
                        'color': 'white', 'cursor': 'pointer',
                    },
                ),
            ], md=6),
            dbc.Col([
                dbc.Button([
                    html.Span("📑", style={'fontSize': '1.4rem', 'marginRight': '8px',
                                           'verticalAlign': 'middle'}),
                    html.Span([
                        html.Strong("Detail Insight Analytics Report",
                                    style={'display': 'block', 'fontSize': '1rem'}),
                        html.Small("5-Table Deep-Dive Analysis  (Word .docx)",
                                   style={'color': '#d4f5e9', 'fontWeight': 'normal',
                                          'fontSize': '0.82rem'}),
                    ], style={'display': 'inline-block', 'verticalAlign': 'middle'}),
                ],
                    id='rt-detail-report-btn',
                    style={
                        'width': '100%', 'padding': '14px 20px',
                        'borderRadius': '10px', 'textAlign': 'left',
                        'background': 'linear-gradient(135deg,#145a32,#1e8449)',
                        'border': 'none', 'boxShadow': '0 3px 10px rgba(30,132,73,0.4)',
                        'color': 'white', 'cursor': 'pointer',
                    },
                ),
            ], md=6),
        ], className='g-3'),
        html.Div(id='rt-report-status', style={'marginTop': '10px'}),
    ], style={
        'background': 'white', 'border': '1px solid #d0d7de',
        'borderRadius': '10px', 'padding': '22px 26px', 'marginBottom': '28px',
        'boxShadow': '0 2px 8px rgba(0,0,0,0.06)',
    }),

    # ══════════════════════════════════════════════════════════════════════
    # SECTION 2 — TRACKING
    # ══════════════════════════════════════════════════════════════════════
    html.Div([
        html.H4("📌 Tracking — Monitor Performance",
                style={'color': '#1a3c5e', 'borderBottom': '2px solid #2d6a9f',
                       'paddingBottom': '8px', 'marginBottom': '16px'}),
        html.P("Select any combination of District, Office and Service, "
               "then click Track to add a row to the watchlist.",
               style={'color': '#666', 'fontSize': '0.92rem', 'marginBottom': '18px'}),

        dbc.Row([
            dbc.Col([
                html.Label("🏛️ District", style={'fontWeight': 'bold', 'marginBottom': '4px'}),
                dcc.Dropdown(id='rt-track-district', options=_district_opts,
                             value=None, placeholder='Select district…', clearable=True),
            ], md=4),
            dbc.Col([
                html.Label("🏢 Office", style={'fontWeight': 'bold', 'marginBottom': '4px'}),
                dcc.Dropdown(id='rt-track-office', options=_office_opts,
                             value=None, placeholder='Select office…', clearable=True),
            ], md=4),
            dbc.Col([
                html.Label("⚙️ Service", style={'fontWeight': 'bold', 'marginBottom': '4px'}),
                dcc.Dropdown(id='rt-track-service', options=_service_opts,
                             value=None, placeholder='Select service…', clearable=True),
            ], md=3),
            dbc.Col([
                html.Label("\u00a0", style={'display': 'block', 'marginBottom': '4px'}),
                dbc.Button("➕ Track", id='rt-track-btn', color='success',
                           style={'width': '100%', 'fontWeight': 'bold'}),
            ], md=1),
        ], className='mb-3 align-items-end'),

        html.Div(id='rt-track-status', style={'marginBottom': '12px'}),

        html.Div([
            html.H6("📋 Tracking Watchlist",
                    style={'color': '#1a3c5e', 'fontWeight': 'bold', 'marginBottom': '10px'}),
            html.Div(id='rt-tracked-table'),
        ], style={'background': '#f8fbff', 'border': '1px solid #d0e4f4',
                  'borderRadius': '8px', 'padding': '14px 18px'}),

    ], style={
        'background': 'white', 'border': '1px solid #d0d7de',
        'borderRadius': '10px', 'padding': '22px 26px',
        'boxShadow': '0 2px 8px rgba(0,0,0,0.06)',
    }),

    # Persistent stores
    dcc.Store(id='rt-tracked-store', storage_type='local', data=[]),
    dcc.Store(id='rt-uploads-store', storage_type='local', data=[]),

    # Download components
    dcc.Download(id='rt-file-download'),
    dcc.Download(id='rt-ai-report-download'),
    dcc.Download(id='rt-detail-report-download'),

], style={'padding': '20px'})


# ═════════════════════════════════════════════════════════════════════════════
# CALLBACKS
# ═════════════════════════════════════════════════════════════════════════════

# ── 1. Save uploaded files + store metadata with remarks ────────────────────
@app.callback(
    Output('rt-uploads-store', 'data'),
    Output('rt-upload-status', 'children'),
    Output('rt-upload-remark', 'value'),
    Input('rt-upload', 'contents'),
    State('rt-upload', 'filename'),
    State('rt-upload-remark', 'value'),
    State('rt-uploads-store', 'data'),
    prevent_initial_call=True,
)
def handle_upload(contents_list, filenames, remark, current_uploads):
    if not contents_list:
        return current_uploads, dash.no_update, dash.no_update

    existing_names = {e['filename'] for e in (current_uploads or [])}
    new_entries    = list(current_uploads or [])
    alerts         = []
    remark_text    = (remark or '').strip()

    for content, name in zip(contents_list, filenames):
        ext  = name.rsplit('.', 1)[-1].upper() if '.' in name else '?'
        icon = _ICON_MAP.get(ext, '📎')
        try:
            _header, b64 = content.split(',', 1)
            file_bytes   = base64.b64decode(b64)
            size_kb      = round(len(file_bytes) / 1024, 1)
            size_str     = f"{size_kb} KB" if size_kb < 1024 else f"{size_kb / 1024:.1f} MB"

            ts        = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_name = "".join(c if (c.isalnum() or c in '._- ()') else '_' for c in name)
            save_name = f"{ts}_{safe_name}"
            with open(os.path.join(_UPLOAD_DIR, save_name), 'wb') as fh:
                fh.write(file_bytes)

            upload_dt = datetime.datetime.now().strftime('%d-%b-%Y %H:%M')

            if name not in existing_names:
                new_entries.append({
                    'filename': name,
                    'saved_as': save_name,
                    'ext':      ext,
                    'size':     size_str,
                    'uploaded': upload_dt,
                    'remark':   remark_text,
                })
                existing_names.add(name)

            alerts.append(
                dbc.Alert([
                    html.Span(f"{icon} ", style={'fontSize': '1.2rem'}),
                    html.Strong(name),
                    html.Span(f"  ({ext}, {size_str})",
                              style={'fontSize': '0.85rem', 'color': '#555', 'marginLeft': '6px'}),
                    html.Span(" ✅ Saved", style={'color': '#27ae60', 'marginLeft': '10px',
                                                   'fontWeight': '600', 'fontSize': '0.9rem'}),
                ], color='success', className='py-2 px-3 mb-1',
                   style={'borderRadius': '6px', 'fontSize': '0.95rem'}),
            )
        except Exception as ex:
            alerts.append(
                dbc.Alert(f"❌ Error saving {name}: {ex}", color='danger', className='py-2')
            )

    # Clear remark textarea after upload
    return new_entries, html.Div(alerts), ''


# ── 2. Render saved-files table (with Remarks column) ───────────────────────
@app.callback(
    Output('rt-uploads-list', 'children'),
    Input('rt-uploads-store', 'data'),
)
def render_uploads(uploads):
    if not uploads:
        return html.P("No files uploaded yet.",
                      style={'color': '#888', 'fontStyle': 'italic'})

    rows = []
    for i, e in enumerate(uploads, 1):
        icon     = _ICON_MAP.get(e.get('ext', ''), '📎')
        saved_as = e.get('saved_as', '')
        remark   = e.get('remark', '') or '—'
        dl_btn = dbc.Button(
            "⬇️ Download", size='sm', color='primary', outline=True,
            id={'type': 'rt-dl-btn', 'index': saved_as},
            style={'padding': '2px 10px', 'fontSize': '0.82rem'},
            disabled=not bool(saved_as),
        )
        rows.append(html.Tr([
            html.Td(str(i),   style={'color': '#888', 'width': '36px', 'textAlign': 'center'}),
            html.Td([html.Span(icon + " "), html.Strong(e.get('filename', '—'))]),
            html.Td(e.get('ext', '—'),      style={'color': '#2d6a9f', 'textAlign': 'center'}),
            html.Td(e.get('size', '—'),     style={'color': '#555', 'textAlign': 'right'}),
            html.Td(e.get('uploaded', '—'), style={'color': '#555', 'fontSize': '0.85rem'}),
            html.Td(remark,                 style={'color': '#444', 'fontSize': '0.85rem',
                                                   'maxWidth': '220px', 'wordBreak': 'break-word'}),
            html.Td(dl_btn,                 style={'textAlign': 'center'}),
        ]))

    return dbc.Table(
        [
            html.Thead(html.Tr([
                html.Th("#"),
                html.Th("File Name"),
                html.Th("Type",        style={'textAlign': 'center'}),
                html.Th("Size",        style={'textAlign': 'right'}),
                html.Th("Uploaded On"),
                html.Th("Remarks"),
                html.Th("Download",    style={'textAlign': 'center'}),
            ]), style={'background': '#eef4fb'}),
            html.Tbody(rows),
        ],
        bordered=True, hover=True, responsive=True, size='sm',
        style={'marginBottom': '0', 'fontSize': '0.9rem'},
    )


# ── 3. Download an uploaded file from disk ───────────────────────────────────
@app.callback(
    Output('rt-file-download', 'data'),
    Input({'type': 'rt-dl-btn', 'index': ALL}, 'n_clicks'),
    State('rt-uploads-store', 'data'),
    prevent_initial_call=True,
)
def download_uploaded_file(n_clicks_list, uploads):
    ctx = callback_context
    if not ctx.triggered or not any(n for n in (n_clicks_list or []) if n):
        return dash.no_update

    triggered_prop = ctx.triggered[0]['prop_id']
    try:
        id_dict  = json.loads(triggered_prop.rsplit('.', 1)[0])
        saved_as = id_dict['index']
    except Exception:
        return dash.no_update

    orig_name = saved_as
    for e in (uploads or []):
        if e.get('saved_as') == saved_as:
            orig_name = e.get('filename', saved_as)
            break

    file_path = os.path.join(_UPLOAD_DIR, saved_as)
    if not os.path.exists(file_path):
        return dash.no_update

    with open(file_path, 'rb') as fh:
        content = fh.read()

    return dcc.send_bytes(content, filename=orig_name)


# ── 4. Add item to tracking store ────────────────────────────────────────────
@app.callback(
    Output('rt-tracked-store', 'data'),
    Output('rt-track-status', 'children'),
    Input('rt-track-btn', 'n_clicks'),
    State('rt-track-district', 'value'),
    State('rt-track-office',   'value'),
    State('rt-track-service',  'value'),
    State('rt-tracked-store',  'data'),
    prevent_initial_call=True,
)
def handle_track(n_clicks, district, office, service, tracked):
    if not any([district, office, service]):
        return tracked, dbc.Alert(
            "⚠️ Please select at least one of District, Office or Service.",
            color='warning', duration=3000, className='py-2',
        )
    entry = {'district': district or '', 'office': office or '', 'service': service or ''}
    key   = _entry_key(entry)
    if key in [_entry_key(e) for e in tracked]:
        return tracked, dbc.Alert(
            "ℹ️ This combination is already in the watchlist.",
            color='info', duration=3000, className='py-2',
        )
    return tracked + [entry], dbc.Alert(
        "✅ Added to tracking watchlist.", color='success', duration=2000, className='py-2',
    )


# ── 5. Remove row from tracking store ────────────────────────────────────────
@app.callback(
    Output('rt-tracked-store', 'data', allow_duplicate=True),
    Input({'type': 'rt-untrack-row', 'index': ALL}, 'n_clicks'),
    State('rt-tracked-store', 'data'),
    prevent_initial_call=True,
)
def handle_untrack_row(n_clicks_list, tracked):
    ctx = callback_context
    if not ctx.triggered or not any(n for n in (n_clicks_list or []) if n):
        return tracked
    triggered_prop = ctx.triggered[0]['prop_id']
    try:
        id_dict = json.loads(triggered_prop.rsplit('.', 1)[0])
        key     = id_dict['index']
    except Exception:
        return tracked
    return [e for e in tracked if _entry_key(e) != key]


# ── 6. Render tracking watchlist table ───────────────────────────────────────
@app.callback(
    Output('rt-tracked-table', 'children'),
    Input('rt-tracked-store', 'data'),
)
def render_tracked_table(tracked):
    if not tracked:
        return html.P(
            "No items tracked yet. Use the dropdowns above to add a combination.",
            style={'color': '#888', 'fontStyle': 'italic'},
        )
    rows = []
    for i, e in enumerate(tracked, 1):
        district = e.get('district') or '—'
        office   = e.get('office')   or '—'
        service  = e.get('service')  or '—'
        key      = _entry_key(e)
        rows.append(html.Tr([
            html.Td(str(i), style={'color': '#888', 'width': '36px', 'textAlign': 'center'}),
            html.Td(district),
            html.Td(office),
            html.Td(service),
            html.Td(dbc.Button(
                "Untrack", size='sm', color='danger', outline=True,
                id={'type': 'rt-untrack-row', 'index': key},
                style={'padding': '2px 12px', 'fontSize': '0.82rem'},
            )),
        ]))
    return dbc.Table(
        [
            html.Thead(html.Tr([
                html.Th("#"), html.Th("District"), html.Th("Office"),
                html.Th("Service"), html.Th("Action"),
            ]), style={'background': '#eef4fb'}),
            html.Tbody(rows),
        ],
        bordered=True, hover=True, responsive=True, size='sm',
        style={'marginBottom': '0', 'fontSize': '0.9rem'},
    )


# ═════════════════════════════════════════════════════════════════════════════
# REPORT HELPERS
# ═════════════════════════════════════════════════════════════════════════════

def _oot_pct(disposed_out, disposed):
    if disposed == 0:
        return 0.0
    return round(disposed_out / disposed * 100, 2)


def _normalise_raw(df_raw):
    df = df_raw.copy()
    df.columns = df.columns.str.strip()
    _cmap = {
        'District_Eng': 'District', 'Service_Eng': 'Service',
        'Office_Eng': 'Office', 'Disposed_Out': 'OOT',
    }
    df.rename(columns={k: v for k, v in _cmap.items() if k in df.columns}, inplace=True)
    for c in ('Received', 'Disposed', 'OOT'):
        if c not in df.columns:
            df[c] = 0
        df[c] = pd.to_numeric(df[c], errors='coerce').fillna(0).astype(int)
    for c in ('District', 'Office', 'Service'):
        if c not in df.columns:
            df[c] = ''
        df[c] = df[c].astype(str).str.strip()
    return df


def _normalise_mt(df_mt):
    dm = df_mt.copy()
    dm.columns = dm.columns.str.strip()
    _mc = {
        'District_name': 'District', 'Office_name': 'Office', 'Service_name': 'Service',
        'application_Disposed_Out_of_time': 'OOT', 'application_Disposed': 'Disposed',
        'application_Received': 'Received',
    }
    dm.rename(columns={k: v for k, v in _mc.items() if k in dm.columns}, inplace=True)
    for c in ('District', 'Office', 'Service'):
        if c in dm.columns:
            dm[c] = dm[c].astype(str).str.strip()
    for c in ('OOT', 'Disposed', 'Received'):
        if c in dm.columns:
            dm[c] = pd.to_numeric(dm[c], errors='coerce').fillna(0).astype(int)
    return dm


def _set_cell_bg(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _docx_make_table(doc, headers, rows_data, header_bg='1A3C5E', alt_bg='EEF4FB'):
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        para = hdr_cells[i].paragraphs[0]
        if not para.runs:
            para.add_run(str(h))
            hdr_cells[i].text = ''
            para.runs[0].text = str(h)
        para.runs[0].bold = True
        para.runs[0].font.size = Pt(8)
        para.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_bg(hdr_cells[i], header_bg)
    for ridx, row_vals in enumerate(rows_data):
        row = t.add_row().cells
        bg  = alt_bg if ridx % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_vals):
            row[ci].text = str(val)
            para = row[ci].paragraphs[0]
            if para.runs:
                para.runs[0].font.size = Pt(8)
            _set_cell_bg(row[ci], bg)
    return t


# ─── AI Report (Digital Gujarat style) ──────────────────────────────────────
def _build_ai_report_docx(fy):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    fy_key   = fy or '2526'
    data     = FY_DATA.get(fy_key, FY_DATA['2526'])
    fy_label = data['label']
    df       = _normalise_raw(data['df'])

    total_disposed  = int(df['Disposed'].sum())
    total_oot       = int(df['OOT'].sum())
    overall_oot_pct = _oot_pct(total_oot, total_disposed)

    dist_grp = df.groupby('District', as_index=False).agg(
        Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
    )
    dist_grp['OOT_Pct']     = dist_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
    dist_grp['State_Share'] = dist_grp['OOT'].apply(
        lambda x: round(x / total_oot * 100, 2) if total_oot else 0)
    worst5_dist = dist_grp.nlargest(5, 'OOT_Pct').reset_index(drop=True)

    svc_grp = df.groupby('Service', as_index=False).agg(
        Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
    )
    svc_grp['OOT_Pct'] = svc_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
    top5_svc = svc_grp.nlargest(5, 'Received').reset_index(drop=True)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    def _h(text, size=14, color='1A3C5E', bold=True, center=False):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        return p

    def _p(text, size=10, color='333333', bold=False, center=False):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        return p

    # Title
    _h("Digital Gujarat", size=20, center=True)
    _h("Applications Disposed – Out of Time", size=13, color='2D6A9F', center=True)
    _p(f"Financial Year: {fy_label}", size=10, color='666666', center=True)
    _p(f"Generated on: {datetime.datetime.now().strftime('%d-%b-%Y %H:%M')}", size=9,
       color='888888', center=True)
    doc.add_paragraph()
    _p(f"In {fy_label}, out of {total_disposed:,} total disposed applications, "
       f"{total_oot:,} ({overall_oot_pct}%) were disposed out of time.",
       size=11, bold=True, color='1A3C5E')
    doc.add_paragraph()

    # Annexure-1: District-wise
    _h("Annexure-1: District-wise Analysis", size=13)
    _p("Top 5 Worst Performing Districts by Out-of-Time %", bold=True)

    _docx_make_table(doc,
        ['Sr.', 'District', 'Received', 'Disposed', 'Out-of-Time', 'OOT %', 'Share in State %'],
        [[str(i+1), r['District'],
          f"{r['Received']:,}", f"{r['Disposed']:,}", f"{r['OOT']:,}",
          f"{r['OOT_Pct']:.2f}%", f"{r['State_Share']:.2f}%"]
         for i, r in worst5_dist.iterrows()],
    )
    doc.add_paragraph()
    _p("Worst Performing Offices within each District:", bold=True)

    for _, d_row in worst5_dist.iterrows():
        dname   = d_row['District']
        d_df    = df[df['District'] == dname]
        off_grp = d_df.groupby('Office', as_index=False).agg(
            Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
        )
        off_grp['OOT_Pct'] = off_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
        worst5_off = off_grp.nlargest(5, 'OOT_Pct').reset_index(drop=True)
        top_oot    = int(worst5_off['OOT'].sum())
        share      = round(top_oot / int(d_row['OOT']) * 100, 2) if d_row['OOT'] > 0 else 0

        doc.add_paragraph()
        _h(f"Bottom 5 Offices in {dname} District ({fy_label})", size=10, color='1A3C5E')
        _p(f"Share of district OOT total = {share:.2f}%", size=9, color='666666')
        _docx_make_table(doc,
            ['District', 'Office', 'Received', 'Disposed', 'Out-of-Time', 'OOT %'],
            [[dname, r['Office'],
              f"{r['Received']:,}", f"{r['Disposed']:,}", f"{r['OOT']:,}",
              f"{r['OOT_Pct']:.2f}%"]
             for _, r in worst5_off.iterrows()],
            header_bg='2D6A9F',
        )
        doc.add_paragraph()

    # Annexure-2: Service-wise
    doc.add_page_break()
    _h("Annexure-2: Service-wise Analysis", size=13)
    _p("Top 5 Services by Total Applications Received", bold=True)
    _docx_make_table(doc,
        ['Sr.', 'Service', 'Received', 'Disposed', 'Out-of-Time', 'OOT %'],
        [[str(i+1), r['Service'],
          f"{r['Received']:,}", f"{r['Disposed']:,}", f"{r['OOT']:,}",
          f"{r['OOT_Pct']:.2f}%"]
         for i, r in top5_svc.iterrows()],
        header_bg='145A32',
    )
    doc.add_paragraph()
    _p("Worst Performing Offices within each Service:", bold=True)

    for _, s_row in top5_svc.iterrows():
        sname   = s_row['Service']
        s_df    = df[df['Service'] == sname]
        off_grp = s_df.groupby('Office', as_index=False).agg(
            Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
        )
        off_grp['OOT_Pct'] = off_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
        worst5_off = off_grp.nlargest(5, 'OOT_Pct').reset_index(drop=True)

        doc.add_paragraph()
        _h(sname, size=10, color='145A32')
        _docx_make_table(doc,
            ['Office', 'Received', 'Disposed', 'Out-of-Time', 'OOT %', 'Share in State %'],
            [[r['Office'],
              f"{r['Received']:,}", f"{r['Disposed']:,}", f"{r['OOT']:,}",
              f"{r['OOT_Pct']:.2f}%",
              f"{round(r['OOT']/total_oot*100,2) if total_oot else 0:.2f}%"]
             for _, r in worst5_off.iterrows()],
            header_bg='145A32',
        )
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Detail Insight Analytics Report ────────────────────────────────────────
def _build_detail_report_docx(fy):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    fy_key   = fy or '2526'
    data     = FY_DATA.get(fy_key, FY_DATA['2526'])
    fy_label = data['label']
    df       = _normalise_raw(data['df'])
    dm       = _normalise_mt(data['df_mt'])

    all_months = sorted(dm['Month_Year'].unique(),
                        key=lambda x: pd.to_datetime(x, format='%b-%Y'))

    # ── Core aggregates ──────────────────────────────────────────────────
    dist_grp = df.groupby('District', as_index=False).agg(
        Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
    )
    dist_grp['OOT_Pct'] = dist_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
    worst5_dist = dist_grp.nlargest(5, 'OOT_Pct').reset_index(drop=True)

    svc_grp = df.groupby('Service', as_index=False).agg(
        Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
    )
    svc_grp['OOT_Pct'] = svc_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
    top5_svc = svc_grp.nlargest(5, 'Received').reset_index(drop=True)

    def _mo_pct(district=None, office=None, service=None):
        mask = pd.Series([True] * len(dm), index=dm.index)
        if district:
            mask &= (dm['District'] == district)
        if office and 'Office' in dm.columns:
            mask &= (dm['Office'] == office)
        if service and 'Service' in dm.columns:
            mask &= (dm['Service'] == service)
        sub = dm[mask].groupby('Month_Year', as_index=False).agg(
            OOT=('OOT', 'sum'), Disposed=('Disposed', 'sum')
        )
        return {r['Month_Year']: _oot_pct(r['OOT'], r['Disposed']) for _, r in sub.iterrows()}

    # ── Doc setup ────────────────────────────────────────────────────────
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    def _h(text, size=13, color='1A3C5E', bold=True, center=False):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        return p

    def _p(text, size=9, color='555555', bold=False, center=False):
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        return p

    # Title
    _h("Digital Seva Setu", size=20, color='1A3C5E', center=True)
    _h("Detail Insight Analytics Report", size=14, color='2D6A9F', center=True)
    _p(f"Financial Year: {fy_label}  |  Generated: {datetime.datetime.now().strftime('%d-%b-%Y %H:%M')}",
       size=10, color='666666', center=True)
    doc.add_paragraph()
    _p("This report provides 5 deep-dive analytical tables covering district, office and "
       "service performance with monthly Out-of-Time % breakdown.", size=10)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════
    # TABLE 1 — Worst 5 Districts → 5 worst offices + monthly OOT%
    # ════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _h("Table 1: Worst 5 Districts by OOT% — Top 5 Worst Offices with Monthly OOT%", size=12)
    _p("Districts ranked by highest OOT%. For each district, the 5 worst offices "
       "are shown with Received, Disposed, OOT and monthly OOT% for each month.", size=9)
    doc.add_paragraph()

    t1_base = ['District', 'Office', 'Received', 'Disposed', 'OOT', 'OOT %']
    t1_hdrs = t1_base + list(all_months)

    for _, d_row in worst5_dist.iterrows():
        dname   = d_row['District']
        _h(f"District: {dname}  |  Overall OOT%: {d_row['OOT_Pct']:.2f}%", size=10)
        d_df    = df[df['District'] == dname]
        off_grp = d_df.groupby('Office', as_index=False).agg(
            Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
        )
        off_grp['OOT_Pct'] = off_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
        worst5_off = off_grp.nlargest(5, 'OOT_Pct')
        rows = []
        for _, o in worst5_off.iterrows():
            mo = _mo_pct(district=dname, office=o['Office'])
            rows.append([dname, o['Office'],
                         f"{o['Received']:,}", f"{o['Disposed']:,}", f"{o['OOT']:,}",
                         f"{o['OOT_Pct']:.2f}%"] +
                        [f"{mo.get(m, 0):.1f}%" for m in all_months])
        _docx_make_table(doc, t1_hdrs, rows, header_bg='1A3C5E')
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════
    # TABLE 2 — Top 5 Services by Received → 5 worst offices + monthly OOT%
    # ════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _h("Table 2: Top 5 Services by Received — Worst 5 Offices with Monthly OOT%", size=12)
    _p("Services ranked by total received. For each service, the 5 worst offices (by OOT%) "
       "are shown with Received, Disposed, OOT and monthly OOT%.", size=9)
    doc.add_paragraph()

    t2_hdrs = ['Service', 'Office', 'Received', 'Disposed', 'OOT', 'OOT %'] + list(all_months)

    for _, s_row in top5_svc.iterrows():
        sname   = s_row['Service']
        _h(f"Service: {sname}  |  Received: {s_row['Received']:,}", size=10, color='145A32')
        s_df    = df[df['Service'] == sname]
        off_grp = s_df.groupby('Office', as_index=False).agg(
            Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
        )
        off_grp['OOT_Pct'] = off_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
        worst5_off = off_grp.nlargest(5, 'OOT_Pct')
        rows = []
        for _, o in worst5_off.iterrows():
            mo = _mo_pct(service=sname, office=o['Office'])
            rows.append([sname, o['Office'],
                         f"{o['Received']:,}", f"{o['Disposed']:,}", f"{o['OOT']:,}",
                         f"{o['OOT_Pct']:.2f}%"] +
                        [f"{mo.get(m, 0):.1f}%" for m in all_months])
        _docx_make_table(doc, t2_hdrs, rows, header_bg='145A32')
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════
    # TABLE 3 — Top 5 Services by Received → 5 worst districts by OOT%
    # ════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _h("Table 3: Top 5 Services by Received — Worst 5 Districts by OOT%", size=12)
    _p("For each top service, the 5 worst districts (by OOT%) are shown.", size=9)
    doc.add_paragraph()

    t3_rows = []
    for _, s_row in top5_svc.iterrows():
        sname = s_row['Service']
        s_df  = df[df['Service'] == sname]
        d_grp = s_df.groupby('District', as_index=False).agg(
            Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
        )
        d_grp['OOT_Pct'] = d_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
        for _, d in d_grp.nlargest(5, 'OOT_Pct').iterrows():
            t3_rows.append([sname, d['District'],
                            f"{d['Received']:,}", f"{d['Disposed']:,}", f"{d['OOT']:,}",
                            f"{d['OOT_Pct']:.2f}%"])

    _docx_make_table(doc,
        ['Service', 'District', 'Received', 'Disposed', 'OOT', 'OOT %'],
        t3_rows, header_bg='7D3C98')
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════
    # TABLE 4 — Worst 5 Districts → Top 5 Services by Received + Service OOT%
    # ════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _h("Table 4: Worst 5 Districts by OOT% — Top 5 Services with Service OOT%", size=12)
    _p("For the 5 worst districts, the top 5 services (by received) are shown "
       "with each service's OOT% within that district.", size=9)
    doc.add_paragraph()

    t4_rows = []
    for _, d_row in worst5_dist.iterrows():
        dname = d_row['District']
        d_df  = df[df['District'] == dname]
        s_grp = d_df.groupby('Service', as_index=False).agg(
            Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
        )
        s_grp['OOT_Pct'] = s_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
        for _, s in s_grp.nlargest(5, 'Received').iterrows():
            t4_rows.append([dname, s['Service'],
                            f"{s['Received']:,}", f"{s['Disposed']:,}", f"{s['OOT']:,}",
                            f"{s['OOT_Pct']:.2f}%"])

    _docx_make_table(doc,
        ['District', 'Service', 'Received', 'Disposed', 'OOT', 'Service OOT %'],
        t4_rows, header_bg='B7770D')
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════
    # TABLE 5 — Top 5 Services → 5 worst offices → monthly Office OOT% + Service OOT%
    # ════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _h("Table 5: Top 5 Services — Worst 5 Offices — Monthly Office OOT% & Service OOT%", size=12)
    _p("For each top service and its 5 worst offices, monthly overall office OOT% "
       "and monthly service-specific OOT% are shown side by side.", size=9)
    doc.add_paragraph()

    # Pair headers: for each month, 2 columns
    month_pair_headers = []
    for m in all_months:
        month_pair_headers += [f"{m} Office OOT%", f"{m} Svc OOT%"]

    t5_base_hdrs = ['Service', 'Office', 'Received', 'Disposed', 'OOT', 'OOT %']
    t5_hdrs = t5_base_hdrs + month_pair_headers

    for _, s_row in top5_svc.iterrows():
        sname   = s_row['Service']
        _h(f"Service: {sname}  |  Received: {s_row['Received']:,}  |  OOT%: {s_row['OOT_Pct']:.2f}%",
           size=10, color='922B21')
        s_df    = df[df['Service'] == sname]
        off_grp = s_df.groupby('Office', as_index=False).agg(
            Received=('Received', 'sum'), Disposed=('Disposed', 'sum'), OOT=('OOT', 'sum')
        )
        off_grp['OOT_Pct'] = off_grp.apply(lambda r: _oot_pct(r['OOT'], r['Disposed']), axis=1)
        worst5_off = off_grp.nlargest(5, 'OOT_Pct')
        rows = []
        for _, o in worst5_off.iterrows():
            off_mo  = _mo_pct(office=o['Office'])
            svc_mo  = _mo_pct(office=o['Office'], service=sname)
            pairs   = []
            for m in all_months:
                pairs += [f"{off_mo.get(m, 0):.1f}%", f"{svc_mo.get(m, 0):.1f}%"]
            rows.append([sname, o['Office'],
                         f"{o['Received']:,}", f"{o['Disposed']:,}", f"{o['OOT']:,}",
                         f"{o['OOT_Pct']:.2f}%"] + pairs)
        _docx_make_table(doc, t5_hdrs, rows, header_bg='922B21')
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── 7. Generate AI Report callback ──────────────────────────────────────────
@app.callback(
    Output('rt-ai-report-download', 'data'),
    Output('rt-report-status', 'children', allow_duplicate=True),
    Input('rt-ai-report-btn', 'n_clicks'),
    State('fy-store', 'data'),
    prevent_initial_call=True,
)
def generate_ai_report(n_clicks, fy):
    try:
        docx_bytes = _build_ai_report_docx(fy)
        fname  = f"AI_Report_{(fy or '2526')}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        status = dbc.Alert("✅ AI Report generated — downloading now…",
                           color='success', duration=4000, className='py-2 mt-2')
        return dcc.send_bytes(docx_bytes, filename=fname), status
    except Exception as ex:
        return dash.no_update, dbc.Alert(f"❌ Error: {ex}", color='danger', className='py-2 mt-2')


# ── 8. Generate Detail Insight Report callback ──────────────────────────────
@app.callback(
    Output('rt-detail-report-download', 'data'),
    Output('rt-report-status', 'children', allow_duplicate=True),
    Input('rt-detail-report-btn', 'n_clicks'),
    State('fy-store', 'data'),
    prevent_initial_call=True,
)
def generate_detail_report(n_clicks, fy):
    try:
        docx_bytes = _build_detail_report_docx(fy)
        fname  = f"Detail_Insight_Report_{(fy or '2526')}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        status = dbc.Alert("✅ Detail Insight Report generated — downloading now…",
                           color='success', duration=4000, className='py-2 mt-2')
        return dcc.send_bytes(docx_bytes, filename=fname), status
    except Exception as ex:
        return dash.no_update, dbc.Alert(f"❌ Error: {ex}", color='danger', className='py-2 mt-2')

